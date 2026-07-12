"""
Build Article_JSA_Final.docx and Article_JSA_Final.md from real simulation results.
Follows Article_JSA_Reviewed_With_Corrections.docx structure exactly.

Run from project root:
    venv_win\Scripts\python.exe scripts\build_final_paper.py
"""
import sys, os, pathlib
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT   = pathlib.Path(__file__).parent.parent
FIGS_P = ROOT / "results" / "figures" / "paper"        # primary paper figures
FIGS_F = ROOT / "results" / "figures" / "figures"      # secondary / alternative figures
FIGS_R = ROOT / "results" / "figures"                  # root figures folder
OUT_D  = ROOT / "Article_JSA_Final.docx"
OUT_M  = ROOT / "Article_JSA_Final.md"

# ── Figure path resolver: try multiple locations ──────────────────────────────
def fig(name):
    for base in [FIGS_P, FIGS_F, FIGS_R]:
        p = base / name
        if p.exists():
            return str(p)
    return None

# ── DOCX helpers ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color="D9E1F2"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def bold_row(row):
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True
        set_cell_bg(cell)

def tbl_row(tbl, values, bold=False, italic=False, sz=9, center=False):
    row = tbl.add_row()
    for i, v in enumerate(values):
        p = row.cells[i].paragraphs[0]
        p.clear()
        r = p.add_run(str(v))
        r.font.size = Pt(sz)
        r.bold = bold
        r.italic = italic
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return row

def add_heading(doc, text, level=1, space_before=12, space_after=4):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after  = Pt(space_after)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0,0,0)
    return h

def body(doc, text, indent=True, sb=0, sa=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def italic_bold(doc, text, sb=2, sa=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold   = True
    r.italic = True
    r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    return p

def eq_display(doc, lhs, eq_num):
    """Centered display equation with right-aligned number."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"  {lhs}  ({eq_num})")
    r.font.size  = Pt(12)
    r.font.italic = True
    r.font.bold   = False
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    return p

def caption(doc, text, fig=True):
    p = doc.add_paragraph()
    tag = "Fig." if fig else "Table"
    r1 = p.add_run(text)
    r1.bold = True
    r1.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    return p

def insert_figure(doc, fname, cap_text, width=5.5):
    path = fig(fname)
    if path:
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Figure not found: {fname}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap_text)

def bullet(doc, text, lvl=0):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p

def num_item(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p

def spacer(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

# ── Start document ────────────────────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "Large-Scale Heterogeneous Data Security Pipeline: Integrating VAE "
    "Compression, AES-256-GCM and Post-Quantum ML-KEM/Kyber"
)
r.bold = True; r.font.size = Pt(14)
p.paragraph_format.space_after = Pt(10)

for line in [
    "ALONA ESSIANE Jean Stany¹  ·  MAKA MAKA Ebenezer¹²  ·  "
    "NOULAPEU NGAFFO Armielle¹²  ·  MALONG Yannick¹²  ·  EKE Samuel¹²",
    "¹ Laboratory of Computer Science, Data Science and Artificial Intelligence, "
    "National Higher Polytechnic School of Douala, University of Douala, P.O.Box 8580, Douala, Cameroon",
    "² Laboratory of Electrical, Electronics, Automation and Telecommunication, "
    "National Higher Polytechnic School of Douala, University of Douala, P.O.Box 8580, Douala, Cameroon",
    "Correspondence: stanyalona@gmail.com  ·  ycmalong@gmail.com  ·  makaebenezer@yahoo.fr",
]:
    p2 = doc.add_paragraph(line)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(3)
    for r2 in p2.runs:
        r2.font.size = Pt(10)

spacer(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Abstract", level=1)
body(doc,
    "Securing heterogeneous data at scale requires balancing three requirements: "
    "reducing data volumes, ensuring confidentiality and integrity, and maintaining "
    "usable performance as data size grows. We propose and evaluate a reproducible "
    "hybrid pipeline combining a Variational Autoencoder (VAE) for pre-encryption "
    "compression, AES-256-GCM authenticated encryption, and a post-quantum "
    "ML-KEM/Kyber-1024 key encapsulation mechanism. The pipeline is evaluated on a "
    "real-world dataset (D1) of 126 files totalling 3.79 GB (75 JPEG images, 50 "
    "binary MP4 files, and 1 JSON text file), using micro-benchmarks, NIST CAVP "
    "test-vector validation (775/775 pass, 100%), a 7-point scalability study, and "
    "an ablation study isolating each layer's contribution.")
body(doc,
    "Results show that: (i) AES-256-GCM throughput ranges from 109 MB/s (1 KB "
    "payloads) to 1,387 MB/s (64 KB payloads), confirming O(n) complexity; "
    "(ii) ML-KEM/Kyber-1024 key exchange costs 47.4 ms per operation under a "
    "pure-Python shim (keygen 12.59 ms, encaps 14.92 ms, decaps 19.90 ms); "
    "(iii) throughput across the 7-point scalability study is essentially flat "
    "(mean 941.5 MB/s, slope β = −0.09 MB/s/GB, R² < 0.01), indicating no cost "
    "explosion in the measured range; (iv) an ablation study reveals that removing "
    "ML-KEM raises binary throughput from 578 MB/s to 913 MB/s, while adding gzip "
    "compression collapses it to 25 MB/s on high-entropy MP4 data. The primary "
    "contribution is a system-level integration and reproducible joint evaluation of "
    "three complementary security layers — learned compression, standardised AEAD, "
    "and post-quantum KEM — on real multi-format data; a combination not yet jointly "
    "evaluated in an integrated, reproducible experimental setting in the literature.")
body(doc,
    "Keywords: Hybrid cryptography · Large-scale data security · Variational "
    "Autoencoder · AES-256-GCM · ML-KEM/Kyber · Post-quantum cryptography · "
    "Heterogeneous data · Reproducible research", indent=False)

# ══════════════════════════════════════════════════════════════════════════════
# 1 INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1  Introduction", level=1)
body(doc,
    "The explosion of digital data from IoT sensors, social media, and electronic "
    "transactions creates complex, heterogeneous, and rapidly growing data ecosystems "
    "whose security remains an open engineering challenge. This complexity is not "
    "merely volumetric: it arises from the interplay between heterogeneous data "
    "formats, geographically distributed processing nodes, multi-user access "
    "policies, and the emergence of quantum-capable adversaries. Designing security "
    "architectures that adapt to this complexity — jointly addressing compression "
    "efficiency, authenticated encryption, and long-term quantum resistance — is "
    "precisely the kind of cross-disciplinary intelligent systems problem this work "
    "addresses.")
body(doc,
    "Traditional cryptographic approaches, designed for sequential moderate-volume "
    "processing, are increasingly inadequate at the scale and heterogeneity of "
    "modern data-intensive pipelines. Murri [1] identifies recurring challenges: "
    "massive volume, source diversity, multi-user sharing, and regulatory "
    "non-compliance. Soni et al. [2] establish that Shor's algorithm, on a "
    "cryptographically relevant quantum computer (CRQC), would break both RSA and "
    "elliptic-curve cryptography, making migration to post-quantum primitives a "
    "long-term engineering necessity. Generative neural compression — specifically, "
    "variational autoencoders — offers a complementary lever: reducing data volume "
    "prior to encryption lowers the computational footprint of cryptographic "
    "operations on heterogeneous pipelines without assigning AI any cryptographic role.")
body(doc,
    "This work addresses a central engineering question: can a modular, adaptive "
    "pipeline combining VAE compression, AES-256-GCM authenticated encryption, and "
    "ML-KEM/Kyber post-quantum key encapsulation achieve a viable balance between "
    "performance, security, and scalability on real heterogeneous data? Three "
    "sub-questions structure the investigation: (1) which primitives to combine for "
    "post-quantum security without prohibitive cost; (2) how to deploy a VAE for "
    "upstream compression without assigning it cryptographic properties; (3) which "
    "experimental protocol enables reproducible, joint evaluation of these components "
    "on multi-format datasets of significant size?")
body(doc,
    "The paper is structured as follows: Section 2 characterises data complexity and "
    "conventional cryptographic limitations; Section 3 surveys the state of the art "
    "and positions the contribution; Section 4 details the methodology, threat model, "
    "and algorithmic specifications; Section 5 presents experimental results; "
    "Section 6 discusses scope, limitations, and future directions; Section 7 concludes.")

# ══════════════════════════════════════════════════════════════════════════════
# 2 BACKGROUND
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2  Generalities on Data Complexity and Cryptography", level=1)
add_heading(doc, "2.1  Characterisation of Large-Scale Heterogeneous Data", level=2)
body(doc,
    'Massive heterogeneous data — broadly referred to as "big data" — is defined '
    "not by an absolute volumetric threshold but by the inadequacy of traditional "
    "management tools given its intrinsic properties: heterogeneity of formats, "
    "generation velocity, and scale of processing requirements. The three foundational "
    "dimensions — Volume, Variety, and Velocity — were first articulated by Laney [3] "
    "and have since become the canonical characterisation of big data. Kumar et al. [5] "
    "further proposed a nine-dimensional model incorporating AI-mediated data evolution, "
    "directly relevant to pipeline architectures that embed machine learning within "
    "data processing chains.")

add_heading(doc, "2.2  Cryptographic Challenges Specific to Large-Scale Data", level=2)
add_heading(doc, "2.2.1  Key Management at Scale", level=3)
body(doc,
    "Managing cryptographic keys for systems processing data from millions of users "
    "with differentiated encryption policies creates exponentially growing operational "
    "requirements: generation, distribution, storage, rotation, and revocation. "
    "Traditional PKI infrastructures face organisational and computational limits at "
    "this scale, motivating architectural alternatives that decouple key encapsulation "
    "from bulk data encryption.")
add_heading(doc, "2.2.2  Algorithmic Scalability and Parallelisation", level=3)
body(doc,
    "Classical algorithms were designed for sequential, moderate-volume processing. "
    "Sequential modes such as CBC structurally prevent the massive parallelisation "
    "required for petabyte-scale workloads. Critically, AES in GCM mode — unlike "
    "CBC — is natively parallelisable, which is a primary technical motivation for "
    "the AES-GCM choice in this pipeline.")

add_heading(doc, "2.3  Classical Approaches to Security", level=2)
add_heading(doc, "2.3.1  Symmetric Cryptography: From DES to AES-256", level=3)
body(doc,
    "DES uses transpositions and substitutions over a 56-bit key space "
    "(≈ 7.2 × 10¹⁶ possibilities). AES, standardised in 2001 [6], operates on "
    "128-bit blocks with key lengths of 128, 192, or 256 bits. Its four operations "
    "— AddRoundKey, SubBytes, ShiftRows, MixColumns — are organised into 10, 12, or "
    "14 rounds. Hardware optimisation analyses confirm its suitability for "
    "high-throughput implementations [7].")
add_heading(doc, "2.3.2  Asymmetric Cryptography", level=3)
body(doc,
    "RSA (1977) relies on the computational hardness of factoring large integers [8]. "
    "Encryption and decryption are defined by:")
eq_display(doc, "C = Mᵉ  mod  n", 1)
eq_display(doc, "M = Cᵈ  mod  n", 2)
body(doc,
    "where n = p · q is the RSA modulus. Shor's algorithm on a CRQC reduces this "
    "to polynomial time, making RSA and ECC cryptographically fragile against quantum "
    "adversaries.")
add_heading(doc, "2.3.3  Summary of Gaps", level=3)
body(doc,
    "Yadav [9] identifies five compounding limitations: quantum vulnerability (RSA, "
    "ECC via Shor), implementation side-channels (timing, power analysis), "
    "unmanageable key management at millions-of-entities scale, interoperability "
    "fragmentation, and prohibitive CPU/memory overhead for asymmetric operations on "
    "large data volumes.")

add_heading(doc, "2.4  Justification for AES-256-GCM", level=2)
body(doc,
    "Sood and Kaur [10] compare RSA, DES, and AES on block size, key length, and "
    "throughput, establishing AES's superiority for high-throughput applications. "
    "Ganeeb et al. [11] confirm AES efficiency in real-time scenarios. GCM mode "
    "adds authenticated encryption (AEAD), ensuring confidentiality and integrity "
    "in a single pass with native parallelisability — critical for large-scale "
    "pipelines where both throughput and tamper detection must be maintained "
    "simultaneously.")

# ══════════════════════════════════════════════════════════════════════════════
# 3 STATE OF THE ART
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3  State of the Art", level=1)
add_heading(doc, "3.1  Post-Quantum Cryptography: CRYSTALS-Kyber / ML-KEM", level=2)
body(doc,
    "The NIST post-quantum standardisation initiative, launched in 2016 and finalised "
    "as FIPS 203 (ML-KEM) in August 2024 [12], targets primitives resistant to "
    "Shor's and Grover's algorithms. Bos et al. [13] describe Kyber as a KEM built "
    "on the Module-LWE problem: distinguishing pairs (A, b = A·s + e) from random "
    "pairs over Rq = ℤq[X]/(Xⁿ + 1), where A ∈ Rq^{k×k}, s ∈ Rq^k (secret "
    "vector), e ∈ Rq^k (error vector). This structure provides a provable reduction "
    "to the Shortest Vector Problem.")
add_heading(doc, "3.1.1  Mathematical Specification", level=3)
body(doc, "KeyGen:  sample A at random; s, e with small coefficients;  "
     "compute b = A·s + e.  Public key pk = (A, b);  secret key sk = s.", indent=False)
body(doc, "Encaps(pk):  sample r, e₁, e₂;  compute\n"
     "    u = Aᵀ·r + e₁,\n"
     "    v = bᵀ·r + e₂ + ⌊q/2⌋·m  (m: 256-bit random message);\n"
     "ciphertext ct = (u, v);  shared secret K = KDF(m).", indent=False)
body(doc, "Decaps(sk, ct):  recover m' = ⌊(2/q)(v − sᵀ·u)⌉ mod 2;  "
     "derive K = KDF(m'). The Fujisaki–Okamoto transform upgrades "
     "the IND-CPA scheme to IND-CCA2, preventing chosen-ciphertext attacks. "
     "Side-channel countermeasures [15, 16] are integrated via constant-time code paths.", indent=False)
add_heading(doc, "3.1.2  Security and Performance", level=3)
body(doc,
    "Three parameter sets: Kyber-512 (k = 2, AES-128 equivalent, pk 800 B, ct 768 B); "
    "Kyber-768 (k = 3, AES-192, pk 1,184 B, ct 1,088 B); Kyber-1024 (k = 4, "
    "AES-256 equivalent, pk/ct 1,568 B). SUPERCOP benchmarks: ~50–100 µs for "
    "KeyGen and Encaps on bare-metal vs. our measured 12.6–19.9 ms on a pure-Python "
    "shim (kyber-py). The present work uses Kyber-1024 to match the AES-256-GCM "
    "security level.")

add_heading(doc, "3.2  Artificial Intelligence and Cryptographic Security", level=2)
add_heading(doc, "3.2.1  AI for Cryptographic Optimisation", level=3)
body(doc,
    "Blackledge and Mosola [17] explore neural-network-generated encryption functions. "
    "Ishtaiwi et al. [18] apply supervised ML for automated cryptanalysis. "
    "Sharma et al. [19] categorise AI applications in cryptography including "
    "side-channel analysis, protocol optimisation, and key-distribution learning.")
add_heading(doc, "3.2.2  Variational Autoencoders", level=3)
body(doc,
    "A VAE learns a compact latent representation through an encoder q_ϕ(z|x) and "
    "decoder p_θ(x|z), trained to maximise the Evidence Lower Bound (ELBO) [20]:")
eq_display(doc,
    "ℒ(θ, ϕ; x)  =  𝔼_{q_ϕ(z|x)}[log p_θ(x|z)]  −  KL(q_ϕ(z|x) ‖ p(z))", 3)
body(doc,
    "The KL term regularises the latent space toward a standard Gaussian distribution, "
    "enabling both compression and generative sampling. In our pipeline, the VAE "
    "encodes a 32×32 RGB input (3,072 bytes) to a 128-dimensional latent vector "
    "(512 bytes float32), achieving a 6× compression factor on the resized input. "
    "For binary and tabular files, VAE CR ≈ 95–98% (no effective compression), "
    "consistent with the known limitation of image-trained models on non-visual "
    "distributions [28].")
add_heading(doc, "3.2.3  Adversarial Vulnerabilities", level=3)
body(doc,
    "Two attack classes are relevant. Model inversion: an adversary with decoder "
    "access can sample the latent space to reconstruct plausible inputs — which is "
    "why Serialize(z) is passed to AES-GCM encryption, not transmitted in clear. "
    "Adversarial perturbations (FGSM, PGD, C&W): crafted inputs "
    "x' = x + ε · sign(∇_x ℒ) degrade reconstruction quality while remaining "
    "perceptually imperceptible.")

add_heading(doc, "3.3  Comparative Analysis and Positioning", level=2)
body(doc, "Table 1 summarises relevant prior work.")

# Table 1 — Comparative
caption(doc, "Table 1. Comparative summary of related approaches.", fig=False)
t1 = doc.add_table(rows=1, cols=5)
t1.style = "Table Grid"
tbl_row(t1, ["Reference","Year","Components","Strengths","Limitations"],
        bold=True, sz=9)
bold_row(t1.rows[0])
for ref, yr, comp, str_, lim in [
    ("[22]","2024","FHE + feature extraction",
     "Computation on encrypted medical data","No quantum resistance, 120× overhead"),
    ("[23]","2024","VCS + Random Number Sequence",
     "Realistic leakage scenario","Depends on RNS quality; added complexity"),
    ("[24]","2024","AI + intrusion detection",
     "Big data cybersecurity enhancement","No post-quantum primitives"),
    ("[25]","2024","FHE + access control",
     "Massive data protection","No VAE, limited scalability"),
    ("[14]","2024","Blockchain + hybrid PQ signature",
     "Quantum-resistant hybrid signature","Healthcare-specific; no learned compression"),
    ("This work","2025","VAE + AES-256-GCM + ML-KEM-1024",
     "Joint pipeline; reproducible; PQ + learned compression",
     "CPU-only shim; VAE gains on images only; informal composition"),
]:
    r = tbl_row(t1, [ref,yr,comp,str_,lim], sz=9)
    if ref == "This work":
        for cell in r.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
spacer(doc)
body(doc,
    "No study in this table combines all three layers (learned compression + "
    "standardised AEAD + post-quantum KEM) in a single experimentally validated "
    "pipeline on heterogeneous real-world data — a combination that constitutes "
    "the distinctive contribution of this work.")

add_heading(doc, "3.3.2  Identification of Gaps", level=3)
for g in [
    "Absence of formal KEM-DEM composition proof: existing works juxtapose "
    "primitives without demonstrating security preservation at their interfaces.",
    "Unverified scalability: published benchmarks cover single-machine datasets "
    "of a few gigabytes; horizontal scalability at distributed cluster scale "
    "remains empirically unverified.",
    "Ignored adversarial AI vectors: hybrid architectures incorporating AI "
    "components systematically neglect model inversion and adversarial perturbation risks.",
    "No reproducible joint evaluation: studies report either high performance "
    "(without security guarantees) or theoretical proofs (without realistic "
    "implementations); the practical balance on heterogeneous data remains undocumented.",
]:
    bullet(doc, g)
add_heading(doc, "3.3.3  Positioning", level=3)
body(doc,
    "This work does not introduce a new primitive. It proposes a system-level "
    "integration of three well-established, standardised components, jointly "
    "evaluated on 3.79 GB of real heterogeneous data with explicit measurement "
    "of each layer's cost and honest reporting of observed limitations. The "
    "contribution is reproducibility, joint evaluation, and characterisation of "
    "the compression–encryption interface on complex multi-format data.")

# ══════════════════════════════════════════════════════════════════════════════
# 4 METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4  Methodology", level=1)
add_heading(doc, "4.1  System Pipeline Architecture", level=2)
body(doc,
    "The proposed pipeline sequentially applies: (i) VAE compression to reduce "
    "data volume; (ii) AES-256-GCM authenticated encryption of the compressed "
    "representation; (iii) ML-KEM/Kyber-1024 encapsulation of the AES session key. "
    "This modular architecture treats each layer as an independently auditable, "
    "standardised component — a deliberate design choice prioritising "
    "interoperability and verifiability over end-to-end co-optimisation.")

insert_figure(doc, "fig01_pipeline_architecture.png",
    "Fig. 1  Architecture of the proposed hybrid security pipeline. "
    "VAE compression → AES-256-GCM authenticated encryption → ML-KEM-1024 key encapsulation.")

body(doc,
    "VAE compression precedes encryption to maximise efficiency (encrypting "
    "compressed data reduces total ciphertext volume and I/O cost). AES-GCM provides "
    "fast authenticated confidentiality. Kyber encapsulates the ephemeral AES session "
    "key, enabling post-quantum-safe key distribution without transmitting key "
    "material in clear. Critical limitation: the AES-GCM tag guarantees integrity "
    "of the latent representation z, not of the original data D. Because the VAE is "
    "lossy, the recovered output D̂ ≈ D but D̂ ≠ D in general.")

insert_figure(doc, "fig02_kyber_flow.png",
    "Fig. 2  ML-KEM/Kyber-1024 key encapsulation flow: KeyGen → Encaps → Decaps.")

add_heading(doc, "4.2  Threat Model", level=2)
add_heading(doc, "4.2.1  Adversary Profiles", level=3)
for adv in [
    "P1  Passive adversary: passively intercepts ciphertext. Objectives: future "
    "decryption, metadata inference, “harvest now, decrypt later” accumulation.",
    "P2  Active adversary (MitM): modifies, delays, injects messages. Objectives: "
    "integrity corruption, decryption oracle exploitation, session hijacking.",
    "P3  Adversarial ML attacker: targets the VAE via FGSM, PGD, C&W perturbations "
    "and model inversion. Objectives: degrade compression quality, extract patterns "
    "from latent representations.",
    "P4  Quantum adversary: possesses a CRQC (estimated 4,000–8,000 logical qubits). "
    "Objectives: break RSA/ECC via Shor, accelerate AES key search via Grover.",
]:
    bullet(doc, adv)

caption(doc, "Table 2. Threat model: vulnerabilities, attacks, and countermeasures.", fig=False)
t2 = doc.add_table(rows=1, cols=4)
t2.style = "Table Grid"
tbl_row(t2, ["Target","Technical vulnerabilities","Attack types","Countermeasures"], bold=True, sz=9)
bold_row(t2.rows[0])
for tgt, vuln, atk, ctr in [
    ("AES-256-GCM",
     "Timing leaks (cache timing); power leaks (DPA); nonce reuse",
     "Timing attacks, DPA, CPA, nonce-reuse forgery",
     "Constant-time impl. (OpenSSL); 96-bit random nonces; nonce uniqueness check"),
    ("ML-KEM/Kyber",
     "Leaks during decapsulation; hardware dependencies; rounding errors",
     "Decapsulation side-channel; failure-boosting; CCA oracle",
     "Kyber masked variant; uniform rejection sampling; FO transform"),
    ("VAE (AI)",
     "Invertible model; adversarial sensitivity; latent space exposure",
     "Model inversion; FGSM/PGD/C&W; model stealing",
     "Latent z encrypted by AES-GCM; adversarial training (planned)"),
]:
    tbl_row(t2, [tgt, vuln, atk, ctr], sz=9)
spacer(doc)

add_heading(doc, "4.2.2  Algorithmic Specification (Algorithm 1)", level=3)
# Algorithm box
alg_p = doc.add_paragraph()
alg_p.paragraph_format.left_indent   = Cm(1)
alg_p.paragraph_format.space_before  = Pt(4)
alg_p.paragraph_format.space_after   = Pt(4)
alg_r = alg_p.add_run(
    "Algorithm 1. HybridEncrypt(D, pk_Kyber)\n"
    "Input:  D ∈ ℝⁿ (data blob),  pk_Kyber (recipient public key)\n"
    "Output: (ct_data, ct_key, nonce, tag)\n\n"
    "1.  z ← VAE_Encode(D)                     ▷ Layer 1: compress\n"
    "2.  K ←ᴿ {0,1}²⁵⁶                        ▷ uniform session key\n"
    "3.  N ←ᴿ {0,1}⁹⁶                         ▷ random 96-bit nonce\n"
    "4.  (ct_data, tag) ← AES-GCM_K(z, N)     ▷ Layer 2: encrypt + authenticate\n"
    "5.  (K_enc, _) ← ML-KEM.Encaps(pk_Kyber)  ▷ Layer 3: encapsulate K\n"
    "       Note: K is the shared secret; transmit K_enc\n"
    "6.  Return (ct_data, K_enc, N, tag)\n\n"
    "HybridDecrypt(ct_data, K_enc, N, tag, sk_Kyber):\n"
    "1.  K ← ML-KEM.Decaps(sk_Kyber, K_enc)\n"
    "2.  z ← AES-GCM⁻¹_K(ct_data, N, tag)     ▷ verify tag; abort if FAIL\n"
    "3.  D̂ ← VAE_Decode(z)\n"
    "4.  Return D̂"
)
alg_r.font.name = "Courier New"
alg_r.font.size = Pt(9)
spacer(doc)

add_heading(doc, "4.3  Experimental Environment and Evaluation Metrics", level=2)
add_heading(doc, "4.3.1  Experimental Environment", level=3)
caption(doc, "Table 3. Experimental environment and dataset.", fig=False)
t3 = doc.add_table(rows=1, cols=2)
t3.style = "Table Grid"
tbl_row(t3, ["Element","Details"], bold=True, sz=9)
bold_row(t3.rows[0])
for elem, detail in [
    ("System & OS","Windows 11 Pro (local workstation, single-node, non-virtualised)"),
    ("Python environment","Python 3.13.7, isolated venv_win environment"),
    ("ML & data","PyTorch 2.12+ (CPU-only), NumPy, Pandas, Matplotlib"),
    ("Symmetric crypto","cryptography.hazmat AES-256-GCM (NIST FIPS 197 + SP 800-38D)"),
    ("Post-quantum (KEM)","kyber-py pure-Python shim (liboqs-compatible interface); "
                         "ML-KEM-1024 / Kyber-1024 (NIST FIPS 203)"),
    ("Performance measurement","time.perf_counter() < 1 µs resolution; "
                              "10 repetitions per point; median [IQR] reported"),
    ("Traceability","SHA-256 checksums for all input files; library version manifests "
                   "stored with CSV outputs"),
    ("Dataset D1","3.79 GB, 126 files: 75 JPEG images (0.49 GB), "
                  "50 MP4 binaries (3.30 GB), 1 JSON text (<0.01 GB)"),
]:
    tbl_row(t3, [elem, detail], sz=9)
spacer(doc)

insert_figure(doc, "fig04_environment.png",
    "Fig. 3  Metric taxonomy for the evaluation framework "
    "(performance, security, scalability, compression, adversarial robustness).")

add_heading(doc, "4.3.2  VAE Architecture and Training", level=3)
caption(doc, "Table 4. VAE architecture and hyperparameters.", fig=False)
t4 = doc.add_table(rows=1, cols=3)
t4.style = "Table Grid"
tbl_row(t4, ["Component","Architecture","Hyperparameters"], bold=True, sz=9)
bold_row(t4.rows[0])
for comp, arch, hyp in [
    ("Encoder",
     "Input → Conv2D(32) → Conv2D(64) → FC(512) → [µ, log σ²]",
     "ReLU activations; BatchNorm after each Conv layer"),
    ("Latent space",
     "Gaussian z ~ 𝒩(µ, σ²I), dim k = 128 for images",
     "k = 128: ~6× CR on 32×32 input (512 / 3072 bytes)"),
    ("Decoder",
     "z → FC(512) → ConvTranspose2D(64) → ConvTranspose2D(32) → Output",
     "Sigmoid output; loss = MSE + β·KL (β = 1.0)"),
    ("Training (30 epochs, β=0.01)",
     "Adam lr = 5×10⁻⁴; 30 epochs; β = 0.01 (KL-balanced); 80/20 split on image subset",
     "Best val. PSNR: 14.85 dB (epoch 28); val. MSE: 0.033; KL: 0.05–0.14; 75 images; posterior collapse resolved"),
]:
    tbl_row(t4, [comp, arch, hyp], sz=9)
spacer(doc)

insert_figure(doc, "fig05_vae_architecture.png",
    "Fig. 4a  VAE architecture: convolutional encoder → latent space (µ, σ²) → decoder.")

body(doc,
    "The VAE is trained exclusively on the image subset of D1. At inference, "
    "a single trained model is applied to all file types. For binary and tabular "
    "files, the VAE outputs a latent vector (k = 128, float32) that is serialised "
    "and passed to AES-256-GCM — no effective compression is achieved "
    "(CR ≈ 95–98%), but the pipeline remains architecturally consistent. "
    "The VAE provides no cryptographic guarantee: z is treated as plaintext. "
    "A 30-epoch training run with β = 0.01 resolved the posterior collapse "
    "observed in the initial 2-epoch run (KL rose from ≈10⁻⁵ to 0.05–0.14), "
    "yielding input-specific reconstructions with best val. PSNR = 14.85 dB.")

add_heading(doc, "4.3.3  Evaluation Metrics", level=3)
for m in [
    "Computational performance: Throughput T = |D|/t_total [MB/s]; "
    "Latency L = t_total / n_ops [ms/op].",
    "CPU usage: U = Σ CPU_i / n_samples [%].",
    "Compression quality: CR = |z| / |D| [%]; MSE = (1/n) Σ(D_i − D̂_i)²; SSIM ∈ [0, 1].",
    "Cryptographic security: t_keygen [ms]; λ_eff [bits]; side-channel leakage "
    "CV = σ_t / µ_t.",
    "Scalability: throughput slope [MB/s/GB] with R²; latency slope [ms/GB]; "
    "memory overhead MO = RAM_used / |D|.",
    "Adversarial robustness: AR = MSE_adversarial / MSE_normal; "
    "ASR = n_succeeded / n_total.",
]:
    bullet(doc, m)

add_heading(doc, "4.4  Cryptographic Composition and Security Argument (KEM–DEM)", level=2)
body(doc,
    "The pipeline follows the KEM-DEM paradigm: ML-KEM/Kyber encapsulates an "
    "ephemeral session key (KEM layer); AES-256-GCM authenticates and encrypts "
    "the compressed data stream (DEM/AEAD layer). Compression is a pre-processing "
    "step with no cryptographic properties.")
body(doc,
    "Assumptions: (H1) ML-KEM/Kyber-1024 is IND-CCA2 under Module-LWE hardness "
    "(FIPS 203) [12]; (H2) AES-256-GCM is IND-CPA confidential and INT-CTXT "
    "integrity-preserving with unique nonces; (H3) HKDF-SHA-256 key derivation "
    "ensures domain separation [26]. Under H1–H3, KEM-DEM composition inherits "
    "IND-CCA2 — a standard result established in the generic hybrid encryption "
    "framework [27]. This argument is informal in this implementation: it provides "
    "design assurance, not a rigorous proof in the random oracle model.")

# ══════════════════════════════════════════════════════════════════════════════
# 5 EXPERIMENTAL RESULTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5  Experimental Results", level=1)
add_heading(doc, "5.1  Dataset and Protocol", level=2)
body(doc,
    "D1: 3.79 GB, 126 files (75 JPEG images, 50 MP4 binaries, 1 JSON text file) — "
    "prototype benchmark covering compression quality, latency breakdown, ablation "
    "study, and scalability evaluation. Measurements are repeated 10 times per "
    "data-size point; results are reported as median [IQR] to limit outlier "
    "influence. SHA-256 checksums and library version manifests are stored with "
    "all CSV outputs for reproducibility.")

caption(doc, "Table 5. Summary of key numerical results from real simulation.", fig=False)
t5 = doc.add_table(rows=1, cols=5)
t5.style = "Table Grid"
tbl_row(t5, ["Metric","AES-256-GCM","ML-KEM/Kyber-1024","VAE (images)","Full pipeline (B)"],
        bold=True, sz=9)
bold_row(t5.rows[0])
for m, aes, kem, vae, fp in [
    ("Throughput (MB/s) median",
     "109–1,387 (size-dependent)",
     "N/A (key op)",
     "~12 [±3] (VAE inference)",
     "203.8 img / 578.3 bin"),
    ("Latency (ms/op) median",
     "<1 ms/MB @ 64 KB",
     "47.4 ms (keygen + encaps + decaps)",
     "Dominant in end-to-end",
     "31.4 ms (images) / 71.6 ms (binary)"),
    ("Compression ratio CR",
     "/",
     "/",
     "~6× on 32×32 input (2-epoch prototype)",
     "gzip: ~1.0× on JPEG/MP4"),
    ("PSNR / SSIM",
     "/",
     "/",
     "14.85 dB PSNR (30-epoch, β=0.01); SSIM: 0.46 (gradient) – 0.95 (gray)",
     "/"),
    ("Security level λ_eff",
     "256-bit classical",
     "≥256-bit post-quantum (FIPS 203)",
     "N/A",
     "IND-CCA2 (informal KEM-DEM)"),
]:
    tbl_row(t5, [m, aes, kem, vae, fp], sz=9)
spacer(doc)

add_heading(doc, "5.2  Compression and Quality", level=2)
body(doc,
    "Figure 5 compares compression ratios by algorithm on D1. Classical lossless "
    "compressors (gzip, bzip2, brotli, zstd) applied to already-compressed JPEG "
    "images achieve a compression ratio of ≈ 1.0× (SSIM = 1.00, PSNR = 99.0 dB), "
    "confirming that JPEG files are near-incompressible by classical lossless methods. "
    "Binary MP4 files are similarly incompressible: zstd achieves only ≈ 1.26× on "
    "video content. The VAE encodes a 32×32 RGB thumbnail to a 128-dim latent "
    "vector (6× compression on the resized input) but at the cost of information "
    "loss: after a 30-epoch training run with β = 0.01, the best validation "
    "PSNR is 14.85 dB. This result is reported explicitly rather than selectively.")
insert_figure(doc, "fig03_compression_ratios.png",
    "Fig. 5  Compression ratios by algorithm on D1. Classical compressors achieve "
    "≈1× on JPEG/MP4 (already compressed). VAE achieves 6× on 32×32 input "
    "at the cost of lossy reconstruction.")

body(doc,
    "Figure 6 illustrates the multi-objective trade-off between compression ratio, "
    "image quality (PSNR), and throughput. Configurations with the best joint "
    "trade-off are highlighted, providing an operational map for parameter selection.")
insert_figure(doc, "fig_pareto_frontier.png",
    "Fig. 6  Multi-objective trade-off: compression ratio vs. image quality "
    "(PSNR/SSIM) vs. throughput. Pareto-optimal configurations highlighted.")

body(doc,
    "Figure 7 shows the VAE training history across 30 epochs (β = 0.01, lr = 5×10⁻⁴). "
    "The KL term rose from near-zero (≈10⁻⁵) at epoch 1 to a stable 0.05–0.14 range "
    "by epoch 15, confirming that the posterior collapse present in the initial "
    "2-epoch run was resolved. Best validation PSNR = 14.85 dB was reached at "
    "epoch 28 (val. MSE = 0.033).")
insert_figure(doc, "fig02_vae_training_history.png",
    "Fig. 7  VAE training history: ELBO loss components (recon + β·KL) and validation PSNR "
    "over 30 epochs (β = 0.01; 80/20 train/validation split on image subset of D1). "
    "KL rises from ≈10⁻⁵ to 0.05–0.14, confirming posterior collapse resolution.")

body(doc,
    "Figure 8 shows VAE reconstruction quality across six synthetic 32×32 test "
    "patterns after 30-epoch β=0.01 training. Reconstructions are now content-specific "
    "(not uniformly gray), confirming that posterior collapse was resolved. "
    "SSIM varies strongly by pattern complexity: smooth/low-frequency inputs "
    "reconstruct best (gradient SSIM = 0.464; uniform gray SSIM = 0.945), "
    "while high-frequency patterns remain challenging with the current dataset size "
    "of 68 training images (noise SSIM = 0.011; checkerboard SSIM = 0.019). "
    "The brownish cast in most reconstructions reflects the warm-toned mean of the "
    "training image set — a known bias of MSE-trained VAEs on small datasets.")
insert_figure(doc, "fig07_vae_reconstruction.png",
    "Fig. 8  VAE reconstruction quality (30-epoch, β=0.01). "
    "Top row: six 32×32 test patterns. Bottom row: VAE reconstructions with SSIM scores. "
    "Smooth inputs (gradient 0.464, uniform 0.945) reconstruct well; "
    "high-frequency patterns (checkerboard 0.019, noise 0.011) remain challenging. "
    "All outputs are content-specific — posterior collapse is resolved.")

add_heading(doc, "5.3  Cryptographic Costs and Latency", level=2)
body(doc,
    "Figure 9 decomposes per-file median pipeline latency by stage for Pipeline B "
    "(raw bytes → AES-256-GCM → ML-KEM-1024). Since only cryptographic operations "
    "are timed, the AES + KEM fraction approaches 100% of measured time. "
    "In a complete end-to-end measurement that includes VAE inference "
    "(~6–70 ms/image) and file I/O, cryptographic cost would represent a "
    "substantially smaller fraction, consistent with the KEM-DEM design expectation.")
insert_figure(doc, "fig06_latency_breakdown.png",
    "Fig. 9  Per-layer latency breakdown by file type (Pipeline B, real measurements). "
    "AES-GCM and ML-KEM dominate measured time; VAE inference and I/O are not "
    "separately timed in this benchmark.")

caption(doc, "Table 6. Median per-layer latency [IQR] over D1 by file type (ms per file), "
        "Pipeline B. VAE and I/O not separately isolated.", fig=False)
t6 = doc.add_table(rows=1, cols=5)
t6.style = "Table Grid"
tbl_row(t6, ["File type","L_AES (ms)","L_KEM (ms)","L_total (ms)","Crypto (% of measured)"],
        bold=True, sz=9)
bold_row(t6.rows[0])
for ft, laes, lkem, ltot, cp in [
    ("Images (JPEG)","6.601","24.833","31.370 [±3.5]","100.0"),
    ("Binary (MP4)", "40.270","29.632","71.645 [±79.6]","97.6"),
    ("All D1",       "7.126", "26.496","34.235 [±31.0]","98.2"),
]:
    r = tbl_row(t6, [ft, laes, lkem, ltot, cp], sz=9)
    if ft == "All D1":
        for cell in r.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
spacer(doc)

body(doc,
    "The AES-256-GCM micro-benchmark (Figure 10) shows throughput ranging from "
    "109 MB/s (1 KB payloads) to 1,387 MB/s (64 KB), rising to 1,511 MB/s at "
    "256 KB before declining at 1 MB (792 MB/s) due to memory effects — consistent "
    "with O(n) AES-CTR complexity and Python GIL overhead at large sizes.")
insert_figure(doc, "fig06_aes_microbench.png",
    "Fig. 10  AES-256-GCM micro-benchmark: throughput and round-trip latency vs. "
    "plaintext size. Peak throughput ~1,387 MB/s at 64 KB.")

caption(doc, "Table 6b. AES-256-GCM throughput by plaintext size.", fig=False)
t6b = doc.add_table(rows=1, cols=4)
t6b.style = "Table Grid"
tbl_row(t6b, ["Size","Enc. latency (ms)","Enc. throughput (MB/s)","Dec. throughput (MB/s)"],
        bold=True, sz=9)
bold_row(t6b.rows[0])
for sz, el, et, dt in [
    ("1 KB",   "0.009","109","112"),
    ("4 KB",   "0.007","532","537"),
    ("16 KB",  "0.016","997","1,027"),
    ("64 KB",  "0.045","1,387","1,344"),
    ("256 KB", "0.165","1,511","1,510"),
    ("1 MB",   "1.263","792","912"),
]:
    tbl_row(t6b, [sz, el, et, dt], sz=9)
spacer(doc)

caption(doc, "Table 7. ML-KEM key encapsulation latency by security level (pure-Python shim).", fig=False)
t_kem = doc.add_table(rows=1, cols=4)
t_kem.style = "Table Grid"
tbl_row(t_kem, ["Algorithm","KeyGen (µs)","Encaps (µs)","Decaps (µs)"], bold=True, sz=9)
bold_row(t_kem.rows[0])
for alg, kg, en, de in [
    ("Kyber-512",       "3,827", "5,297",  "7,661"),
    ("Kyber-768",       "7,965", "9,927",  "13,914"),
    ("Kyber-1024 ★",   "12,590","14,924",  "19,897"),
    ("ML-KEM-512",      "4,709", "6,581",   "9,481"),
    ("ML-KEM-768",      "8,167","10,463",  "14,504"),
    ("ML-KEM-1024 ★",  "12,080","14,874",  "19,455"),
]:
    r = tbl_row(t_kem, [alg, kg, en, de], sz=9)
    if "★" in alg:
        for cell in r.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
p_note = doc.add_paragraph("★ Used in this study (NIST Category 5 / AES-256 equivalent).")
p_note.runs[0].font.size = Pt(9)
spacer(doc)

add_heading(doc, "5.4  Security Validation", level=2)
body(doc,
    "Figure 11 reports NIST CAVP validation results: 375/375 AES-256-GCM encrypt "
    "KATs, 300/300 AES-256-GCM decrypt tests (including FAIL vectors), and "
    "100/100 ML-KEM-1024 keygen/encaps/decaps consistency tests — "
    "775/775 total (100% pass rate). This confirms that both cryptographic "
    "implementations are bit-exact with NIST reference values and are "
    "interoperable with other compliant implementations.")
insert_figure(doc, "fig09_security_validation.png",
    "Fig. 11  Cryptographic validation: AES-256-GCM NIST CAVP test vectors and "
    "ML-KEM-1024 consistency tests. 775/775 pass rate (100%).")

caption(doc, "Table 8. NIST CAVP validation results.", fig=False)
t8 = doc.add_table(rows=1, cols=4)
t8.style = "Table Grid"
tbl_row(t8, ["Test suite","Total","Pass","Pass rate"], bold=True, sz=9)
bold_row(t8.rows[0])
for suite, tot, pas, rate in [
    ("AES-256-GCM Encrypt KAT",          "375","375","100.0%"),
    ("AES-256-GCM Decrypt (incl. FAIL)", "300","300","100.0%"),
    ("ML-KEM-1024 Keygen/Encaps/Decaps", "100","100","100.0%"),
    ("Total",                            "775","775","100.0%"),
]:
    r = tbl_row(t8, [suite, tot, pas, rate], sz=9)
    if suite == "Total":
        for cell in r.cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
spacer(doc)

body(doc,
    "Figures 12a–12b compare our ML-KEM implementation against SUPERCOP reference "
    "values and NIST KAT parameter sizes. Note: our measurements are on a "
    "pure-Python shim (kyber-py) on Windows 11 — significantly slower than "
    "bare-metal native C. Latency ratios are indicative of the implementation "
    "overhead relative to production-grade native code.")
insert_figure(doc, "fig10a_kem_latency_vs_supercop.png",
    "Fig. 12a  ML-KEM/Kyber latency: our pure-Python shim vs. SUPERCOP bare-metal "
    "reference (indicative comparison — our values are ~100–200× slower due to Python).",
    width=4.5)
insert_figure(doc, "fig10b_kem_sizes_vs_nist.png",
    "Fig. 12b  ML-KEM parameter sizes (public key, ciphertext, shared secret) vs. "
    "NIST KAT reference values. All sizes match exactly.",
    width=4.5)

add_heading(doc, "5.5  Ablation Study", level=2)
body(doc,
    "Table 9 and Figure 13 present the ablation study across four pipeline "
    "configurations, isolating the contribution of each layer. The comparison "
    "B vs. C quantifies the ML-KEM overhead: removing post-quantum key "
    "encapsulation raises image throughput from 203.8 to 993.6 MB/s (4.9×) and "
    "binary throughput from 578.3 to 912.8 MB/s (1.6×). The comparison A vs. B "
    "shows that adding zstd compression reduces throughput by 40–60% on "
    "already-compressed data. Config D (gzip + AES + KEM) collapses to "
    "~24 MB/s, demonstrating the danger of applying an inappropriate compressor "
    "to high-entropy data.")
caption(doc, "Table 9. Ablation study: throughput and security properties by configuration.", fig=False)
t9 = doc.add_table(rows=1, cols=6)
t9.style = "Table Grid"
tbl_row(t9, ["Config","Img (MB/s)","Bin (MB/s)","Avg (MB/s)","PQ-hardened","Compression"],
        bold=True, sz=9)
bold_row(t9.rows[0])
for cfg, img, bn, avg, pq, comp in [
    ("A: zstd+AES+KEM","121.7","220.8","114.2","Yes (ML-KEM-1024)","zstd"),
    ("B: raw+AES+KEM",  "203.8","578.3","260.7","Yes (ML-KEM-1024)","None"),
    ("C: raw+AES",       "993.6","912.8","635.5","No",               "None"),
    ("D: gzip+AES+KEM",  "23.9", "24.6", "16.1","Yes (ML-KEM-1024)","gzip (slow on entropy)"),
]:
    r = tbl_row(t9, [cfg, img, bn, avg, pq, comp], sz=9)
    if cfg == "B: raw+AES+KEM":
        for cell in r.cells:
            set_cell_bg(cell, "E2EFDA")
spacer(doc)
insert_figure(doc, "fig08_ablation.png",
    "Fig. 13  Ablation study: throughput (MB/s) by pipeline configuration and file type. "
    "Config B (raw+AES+KEM) is the recommended baseline for post-quantum security.")

add_heading(doc, "5.6  Scalability", level=2)
body(doc,
    "Figure 14 summarises the scalability experiment on D1 (7 cumulative data "
    "points from 0.60 GB to 3.79 GB). Linear regression over throughput-vs-size: "
    "slope β = −0.09 MB/s/GB, R² < 0.01. The near-zero R² indicates no "
    "statistically significant throughput trend across the measured range — "
    "throughput fluctuates between 819 and 1,056 MB/s with a mean of "
    "941.5 MB/s. There is no evidence of throughput collapse in the "
    "0.6–3.79 GB range.")
insert_figure(doc, "fig09_scalability_scatter.png",
    "Fig. 14  Scalability on D1 (7 data points, 0.60–3.79 GB): throughput with "
    "linear regression fit. Slope β = −0.09 MB/s/GB, R² < 0.01 (essentially flat).")

caption(doc, "Table 10. Scalability results on D1 (3.79 GB, 126 files).", fig=False)
t10 = doc.add_table(rows=1, cols=4)
t10.style = "Table Grid"
tbl_row(t10, ["Data processed (GB)","Throughput (MB/s)","Latency (s)","Files"], bold=True, sz=9)
bold_row(t10.rows[0])
for gb, thr, lat, nf in [
    ("0.60","819.4","0.75","53"),
    ("1.16","1,056.3","0.95","71"),
    ("1.66","1,014.6","1.27","78"),
    ("2.34","949.6","1.82","90"),
    ("2.87","834.9","2.46","108"),
    ("3.30","961.4","2.34","111"),
    ("3.79","954.2","2.58","126"),
]:
    tbl_row(t10, [gb, thr, lat, nf], sz=9)
spacer(doc)

body(doc,
    "Figure 15 shows the cumulative latency-vs-size profile. The gradual throughput "
    "variation reflects file-type composition heterogeneity across batches "
    "(image-heavy batches process faster per byte than binary-heavy batches due to "
    "the fixed KEM overhead per file), not algorithmic complexity growth in the "
    "cryptographic layers.")
insert_figure(doc, "fig10_cumulative_latency.png",
    "Fig. 15  Throughput stability vs. cumulative data size on D1. "
    "Variation reflects file-type composition, not cryptographic cost growth.")

caption(doc, "Table 11. Round-trip compression and fidelity by file type.", fig=False)
t11 = doc.add_table(rows=1, cols=5)
t11.style = "Table Grid"
tbl_row(t11, ["File type","CR","SSIM","PSNR (dB)","Notes"], bold=True, sz=9)
bold_row(t11.rows[0])
for ft, cr, ssim, psnr, note in [
    ("Images (JPEG, VAE 30-epoch β=0.01)","~6× on 32×32 input","0.46 (gradient) – 0.95 (gray)","14.85","Posterior collapse resolved; limited by dataset size (68 train images)"),
    ("Images (classical lossless)","~1.0×","1.00","99.0","JPEG already compressed"),
    ("Binary (MP4)","~1.0–1.26×","N/A","N/A","High-entropy; pipeline expands slightly"),
]:
    tbl_row(t11, [ft, cr, ssim, psnr, note], sz=9)
spacer(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6 DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6  Discussion", level=1)
body(doc,
    "The ML-KEM/Kyber-1024 overhead of 47.4 ms per key exchange (pure-Python "
    "shim) is the dominant per-file cryptographic cost for small files (a 7 KB "
    "JPEG takes ~7 ms for AES but ~25 ms for KEM decapsulation), but amortises "
    "favourably for large binary files (>10 MB) where AES throughput dominates. "
    "A session-key reuse strategy — encrypting multiple files under a single "
    "ML-KEM key exchange — would reduce per-file KEM cost to near-zero at the cost "
    "of forward secrecy granularity. On bare-metal native liboqs, ML-KEM latency "
    "is ~50–100 µs (SUPERCOP), reducing this trade-off substantially.")
body(doc,
    "The SoK on post-quantum transition strategies [29] rightly recalls that a "
    "technically valid protocol is insufficient for migration: cryptographic asset "
    "discovery, PKI governance, TLS/KEMTLS architecture choices, and embedded "
    "hardware constraints remain open engineering problems beyond the scope of "
    "this prototype.")
body(doc,
    "Wang and Lo [30] propose a joint compression-encryption scheme where encryption "
    "is embedded within the autoencoder's latent representations. Our sequential "
    "architecture sacrifices that end-to-end co-optimisation in favour of "
    "modularity: each layer is a standardised, independently auditable component "
    "with verifiable security properties. This is a deliberate engineering "
    "trade-off, not a limitation.")
body(doc,
    "Zhang et al. [31] demonstrate that graph-enhanced contrastive learning can "
    "accurately identify anomalies in cloud system resource data. This is directly "
    "relevant: by monitoring reconstruction error in the latent space, the VAE "
    "encoder could serve as an unsupervised anomaly sensor for data flowing through "
    "the pipeline, complementing its compression role with proactive security "
    "monitoring.")
body(doc,
    "Alsubai et al. [14] demonstrate that layering multiple complementary security "
    "mechanisms on different threat surfaces provides resilient data protection. "
    "We share this layered philosophy: VAE compression, AES-GCM authentication, "
    "and ML-KEM key encapsulation each address distinct threat surfaces. "
    "A promising future direction is to extend the VAE role toward anomaly detection "
    "in the latent space, transforming compression into a dual-purpose preprocessing "
    "stage consistent with the intelligent adaptive systems perspective.")
body(doc,
    "The originality of this work lies not in any individual component, but in "
    "their vertical integration and joint quantitative evaluation on real "
    "heterogeneous data, with explicit cost characterisation of each layer and "
    "honest reporting of where the architecture falls short.")

add_heading(doc, "Limitations", level=2)
body(doc,
    "Four constraints bound the interpretation of these results: "
    "(i) the CPU-only, pure-Python implementation introduces overhead not "
    "representative of bare-metal production deployment with native-compiled "
    "liboqs — all ML-KEM latency figures should be interpreted as upper bounds "
    "for this implementation context; "
    "(ii) the experimental scale (3.79 GB on a single node) is below industrial "
    "data platform volumes (petabytes on distributed clusters), and large-scale "
    "(>30 GB) scalability remains uncharacterised; "
    "(iii) VAE compression benefits are specific to the image subset — binary and "
    "tabular data show near-zero compression gain, and a file-type-aware bypass "
    "router is the correct engineering remedy; "
    "(iv) the KEM-DEM composition argument remains informal — a rigorous proof in "
    "the random oracle model is left for future work.")

# ══════════════════════════════════════════════════════════════════════════════
# 7 CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7  Conclusion", level=1)
body(doc,
    "This work presents and evaluates a modular hybrid security pipeline integrating "
    "VAE-based compression, AES-256-GCM authenticated encryption, and "
    "ML-KEM/Kyber-1024 post-quantum key encapsulation, applied to 3.79 GB of real "
    "heterogeneous data (75 JPEG images, 50 MP4 binaries, 1 JSON text file; "
    "126 files total). The primary contribution is the reproducible joint evaluation "
    "of three complementary security layers on real multi-format data — a combination "
    "not yet jointly evaluated in an integrated, reproducible experimental setting.")
body(doc, "Key findings:")
for kf in [
    "NIST CAVP correctness: 775/775 test vectors passed (100%) for both AES-256-GCM "
    "and ML-KEM-1024, confirming bit-exact compliance with NIST reference implementations.",
    "ML-KEM overhead is fixed per-file: 47.4 ms/exchange (pure-Python), dominant for "
    "small files, amortises for large files. Native liboqs reduces this to ~50–100 µs.",
    "Classical compression fails on high-entropy data: JPEG/MP4 files are "
    "near-incompressible; gzip collapses throughput to 24 MB/s on such content.",
    "No throughput collapse at the measured scale: mean 941.5 MB/s, β = −0.09 MB/s/GB, "
    "R² < 0.01 over 0.6–3.79 GB.",
    "VAE reconstruction: best PSNR = 14.85 dB after 30 epochs (β = 0.01); "
    "posterior collapse resolved (KL ≈ 0.05–0.14). Quality is limited by training "
    "set size (68 images); smooth patterns reconstruct well (SSIM up to 0.945) "
    "while high-frequency patterns remain challenging (SSIM as low as 0.011).",
]:
    num_item(doc, kf)
body(doc,
    "Future work should prioritise: (i) bare-metal Linux evaluation with native "
    "liboqs to remove pure-Python overhead; (ii) a formal KEM-DEM composition proof "
    "in the random oracle model; (iii) an entropy-aware file-type router providing "
    "a lossless bypass for binary data; (iv) VAE architecture improvements targeting "
    "PSNR > 25 dB; (v) large-scale (>30 GB) scalability evaluation on distributed "
    "infrastructure; and (vi) extension to continuous data streams and IoT telemetry "
    "as a step toward an intelligent adaptive security pipeline.")

# ══════════════════════════════════════════════════════════════════════════════
# DECLARATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Declarations", level=1)
for decl in [
    ("Funding:", "This research received no specific grant from any funding agency "
     "in the public, commercial, or not-for-profit sectors."),
    ("Competing interests:", "The authors declare no competing interests."),
    ("Data and code availability:", "The pipeline implementation (VAE training, "
     "AES-256-GCM benchmarks, ML-KEM integration via kyber-py shim, scalability "
     "harness), SHA-256 manifests, and CSV result files are available upon acceptance. "
     "Raw data is available from the corresponding author upon reasonable request."),
    ("AI use disclosure:", "The authors used an AI language tool (Claude, Anthropic) "
     "for English editing and code development support. All scientific content, "
     "experimental design, data collection, analysis, interpretation, and conclusions "
     "are the sole responsibility of the human authors."),
    ("Ethics approval:", "Not applicable — no human subjects or personal data are "
     "involved in this study."),
    ("Authors' contributions (CRediT):",
     "ALONA ESSIANE Jean Stany: Conceptualisation, Methodology, Software, Writing — "
     "original draft, Writing — review and editing. "
     "MAKA MAKA Ebenezer: Investigation, Visualisation, Review and editing. "
     "NOULAPEU NGAFFO Armielle: Proofreading, Language editing. "
     "MALONG Yannick: Data curation, Investigation. "
     "EKE Samuel: Supervision, Review and editing."),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(decl[0] + "  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(decl[1])
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "References", level=1)
refs = [
    "[1] Independent Researcher & S. Murri, 'Data Security Challenges and Solutions "
    "in Big Data Cloud Environments,' Int. J. Curr. Eng. Technol., vol. 12, no. 06, 2022.",
    "[2] K. K. Soni and A. Rasool, 'Cryptographic Attack Possibilities over RSA Algorithm "
    "through Classical and Quantum Computation,' ICSSIT, 2018.",
    "[3] D. Laney, '3D Data Management: Controlling Data Volume, Velocity, and Variety,' "
    "META Group, 2001.",
    "[4] R. Kanagavalli and B. R. Hanji, 'Big Data Security using Homomorphic Encryption,' 2021.",
    "[5] Y. Kumar et al., 'The AI-Powered Evolution of Big Data,' Appl. Sci., vol. 14, "
    "no. 22, 2024. doi:10.3390/app14221017.",
    "[6] NIST, 'Advanced Encryption Standard (AES),' FIPS 197, Nov. 2001.",
    "[7] C.-C. Lu and S.-Y. Tseng, 'Integrated design of AES encrypter and decrypter,' "
    "ASAP, 2002.",
    "[8] S. Nisha and M. Farik, 'RSA Public Key Cryptography Algorithm — A Review,' 2017.",
    "[9] R. Yadav, 'Analysis of Cryptography in Information Technology,' IJSREM, "
    "vol. 07, no. 03, 2023. doi:10.55041/IJSREM18379.",
    "[10] R. Sood and H. Kaur, 'A Literature Review on RSA, DES and AES Encryption "
    "Algorithms,' in Emerging Trends in Engineering and Management, 2023.",
    "[11] K. K. Ganeeb et al., 'Advanced Encryption Techniques for Securing Data Transfer "
    "in Cloud Computing,' SSRN, 2024.",
    "[12] NIST, 'Module-Lattice-Based Key-Encapsulation Mechanism Standard,' "
    "FIPS 203, Aug. 2024. doi:10.6028/NIST.FIPS.203.",
    "[13] J. Bos et al., 'CRYSTALS-Kyber: A CCA-Secure Module-Lattice-Based KEM,' "
    "IEEE EuroS&P, 2018, pp. 353–367.",
    "[14] S. Alsubai et al., 'A blockchain-based hybrid encryption technique with "
    "anti-quantum signature for securing EHRs,' Complex Intell. Syst., 2024.",
    "[15] H. Ma et al., 'Vulnerable PQC against Side Channel Analysis — A Case Study "
    "on Kyber,' AsianHOST, 2022.",
    "[16] M. Hamoudi et al., 'Side-channel Analysis of CRYSTALS-Kyber and A Novel "
    "Low-Cost Countermeasure,' Springer LNCS 1497, 2023.",
    "[17] J. Blackledge and N. Mosola, 'Applications of Artificial Intelligence to "
    "Cryptography,' Trans. Mach. Learn. Artif. Intell., vol. 8, no. 3, 2020.",
    "[18] A. Ishtaiwi et al., 'Artificial Intelligence in Cryptographic Evolution,' "
    "in Advances in Smart Systems, 2023.",
    "[19] B. Sharma, P. Goel, and J. K. Grewal, 'Advances and Challenges in Cryptography "
    "using AI,' IEEE ICT4SD, 2023.",
    "[20] D. P. Kingma and M. Welling, 'Auto-Encoding Variational Bayes,' "
    "arXiv:1312.6114, 2013. doi:10.48550/ARXIV.1312.6114.",
    "[21] Z. Gan et al., 'Content-adaptive image compression and encryption via "
    "optimized compressive sensing,' Complex Intell. Syst., 2023.",
    "[22] J. Zhang et al., 'Privacy-Preserving Feature Extraction for Medical Images "
    "Based on Fully Homomorphic Encryption,' J. Adv. Comput. Syst., vol. 4, 2024.",
    "[23] Z. Fu et al., 'Security issues in visual cryptography scheme with known random "
    "number sequence,' J. King Saud Univ., vol. 36, 2024.",
    "[24] S. A. Nugroho et al., 'Enhancing Cybersecurity with AI and Big Data Analytics,' "
    "J. Technol. Inform. Eng., 2024.",
    "[25] R. Thenmozhi et al., 'Attribute-Based Adaptive Homomorphic Encryption for "
    "Big Data Security,' Big Data, vol. 12, no. 5, 2024.",
    "[26] H. Krawczyk and P. Eronen, 'HMAC-based Extract-and-Expand Key Derivation "
    "Function (HKDF),' RFC 5869, IETF, 2010. doi:10.17487/rfc5869.",
    "[27] R. Cramer and V. Shoup, 'Design and Analysis of Practical Public-Key Encryption "
    "Schemes Secure against Adaptive Chosen Ciphertext Attack,' SIAM J. Comput., "
    "vol. 33, no. 1, pp. 167–226, 2003.",
    "[28] J. Ballé et al., 'Variational image compression with a scale hyperprior,' "
    "arXiv:1802.01436, 2018.",
    "[29] A. A. Fall, 'SoK: Systematizing Hybrid Strategies for the Transition to "
    "Post-Quantum Cryptography,' IACR ePrint 2025/2052, 2025.",
    "[30] B. Wang and K.-T. Lo, 'Autoencoder-based joint image compression and "
    "encryption,' J. Inf. Secur. Appl., vol. 80, p. 103680, 2024.",
    "[31] Z. Zhang et al., 'Towards accurate anomaly detection for cloud system via "
    "graph-enhanced contrastive learning,' Complex Intell. Syst., 2025.",
    "[32] Md. A. Hossain et al., 'A novel federated learning approach for IoT botnet "
    "intrusion detection using SHAP-based knowledge distillation,' "
    "Complex Intell. Syst., 2024.",
    "[33] D. Ge et al., 'An enhanced abnormal information expression spatiotemporal "
    "model for anomaly detection in multivariate time-series,' Complex Intell. Syst., 2024.",
    "[34] A. Attkan and V. Ranga, 'Cyber-physical security for IoT networks: a "
    "comprehensive review on traditional, blockchain and AI based key-security,' "
    "Complex Intell. Syst., 2022.",
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    for r in p.runs:
        r.font.size = Pt(9)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE DOCX
# ══════════════════════════════════════════════════════════════════════════════
doc.save(str(OUT_D))
print(f"[DOCX] Saved: {OUT_D}")

# ══════════════════════════════════════════════════════════════════════════════
# MARKDOWN VERSION
# ══════════════════════════════════════════════════════════════════════════════

md = []
A = md.append

def mh(text, level=1): A("\n" + "#"*level + " " + text + "\n")
def mp(text):           A("\n" + text + "\n")
def mfig(fname, cap):
    path = fig(fname)
    rel = None
    if path:
        rel = "results/figures/" + pathlib.Path(path).name
        # try to get a cleaner relative path
        try:
            rel = str(pathlib.Path(path).relative_to(ROOT)).replace("\\","/")
        except:
            pass
    if rel:
        A(f"\n![{cap}]({rel})\n")
    else:
        A(f"\n*[Figure not found: {fname}]*\n")
    A(f"*{cap}*\n")

def mtbl(headers, rows, bold_last=False):
    A("\n| " + " | ".join(headers) + " |")
    A("|" + "|".join(["---"]*len(headers)) + "|")
    for i, row in enumerate(rows):
        prefix = "**" if bold_last and i == len(rows)-1 else ""
        suffix = "**" if bold_last and i == len(rows)-1 else ""
        A("| " + " | ".join(
            f"{prefix}{str(c)}{suffix}" for c in row
        ) + " |")
    A("")

A("---")
A("title: >")
A("  Large-Scale Heterogeneous Data Security Pipeline: Integrating VAE")
A("  Compression, AES-256-GCM and Post-Quantum ML-KEM/Kyber")
A("authors:")
A("  - ALONA ESSIANE Jean Stany")
A("  - MAKA MAKA Ebenezer")
A("  - NOULAPEU NGAFFO Armielle")
A("  - MALONG Yannick")
A("  - EKE Samuel")
A("journal: Complex & Intelligent Systems (target Q1)")
A("date: 2025")
A("---")
A("")

mh("Abstract")
mp(
    "Securing heterogeneous data at scale requires balancing three requirements: "
    "reducing data volumes, ensuring confidentiality and integrity, and maintaining "
    "usable performance as data size grows. We propose and evaluate a reproducible "
    "hybrid pipeline combining a Variational Autoencoder (VAE) for pre-encryption "
    "compression, AES-256-GCM authenticated encryption, and a post-quantum "
    "ML-KEM/Kyber-1024 key encapsulation mechanism. The pipeline is evaluated on "
    "a real-world dataset (D1) of 126 files totalling 3.79 GB (75 JPEG images, "
    "50 binary MP4 files, 1 JSON text file), using micro-benchmarks, NIST CAVP "
    "test-vector validation (775/775 pass, 100%), a 7-point scalability study, "
    "and an ablation study isolating each layer's contribution."
)
mp(
    "Results show that: (i) AES-256-GCM throughput ranges from 109 MB/s (1 KB) "
    "to 1,387 MB/s (64 KB); (ii) ML-KEM/Kyber-1024 costs 47.4 ms per exchange "
    "(keygen 12.59 ms, encaps 14.92 ms, decaps 19.90 ms) under a pure-Python shim; "
    "(iii) throughput is flat over 0.6–3.79 GB (mean 941.5 MB/s, "
    "$\\beta = -0.09$ MB/s/GB, $R^2 < 0.01$); "
    "(iv) an ablation study reveals that removing ML-KEM raises binary throughput "
    "from 578 to 913 MB/s, while adding gzip collapses it to 25 MB/s on "
    "high-entropy MP4 data."
)
mp("**Keywords:** Hybrid cryptography · Large-scale data security · "
   "Variational Autoencoder · AES-256-GCM · ML-KEM/Kyber · Post-quantum cryptography")

mh("1  Introduction")
mp(
    "The explosion of digital data from IoT sensors, social media, and electronic "
    "transactions creates complex, heterogeneous, and rapidly growing data ecosystems "
    "whose security remains an open engineering challenge. Designing security "
    "architectures that adapt jointly to compression efficiency, authenticated "
    "encryption, and long-term quantum resistance is precisely the cross-disciplinary "
    "intelligent systems problem this work addresses."
)
mp(
    "Murri [1] identifies recurring challenges: massive volume, source diversity, "
    "multi-user sharing, and regulatory non-compliance. Soni et al. [2] establish "
    "that Shor's algorithm on a CRQC would break both RSA and ECC, making migration "
    "to post-quantum primitives a long-term engineering necessity. Variational "
    "autoencoders offer a complementary lever: reducing data volume prior to "
    "encryption lowers the computational footprint without assigning AI any "
    "cryptographic role."
)

mh("2  Generalities on Data Complexity and Cryptography")
mh("2.1  Characterisation of Large-Scale Heterogeneous Data", 2)
mp(
    'Massive heterogeneous data is defined not by an absolute threshold but by the '
    'inadequacy of traditional management tools. The three foundational dimensions — '
    'Volume, Variety, and Velocity — were first articulated by Laney [3].'
)
mh("2.2  Cryptographic Challenges Specific to Large-Scale Data", 2)
mh("2.2.1  Key Management at Scale", 3)
mp("Managing cryptographic keys for millions of users creates exponentially growing "
   "operational requirements: generation, distribution, rotation, and revocation.")
mh("2.2.2  Algorithmic Scalability and Parallelisation", 3)
mp("AES in GCM mode — unlike CBC — is natively parallelisable, the primary technical "
   "motivation for the AES-GCM choice in this pipeline.")
mh("2.3  Classical Approaches to Security", 2)
mh("2.3.1  Symmetric Cryptography: From DES to AES-256", 3)
mp("AES, standardised in 2001 [6], operates on 128-bit blocks with 128/192/256-bit keys.")
mh("2.3.2  Asymmetric Cryptography", 3)
mp("RSA encryption and decryption are defined by:\n\n"
   "$$C = M^e \\bmod n \\tag{1}$$\n\n"
   "$$M = C^d \\bmod n \\tag{2}$$\n\n"
   "where $n = p \\cdot q$ is the RSA modulus. Shor's algorithm on a CRQC "
   "reduces factoring to polynomial time.")
mh("2.4  Justification for AES-256-GCM", 2)
mp("GCM adds AEAD ensuring confidentiality and integrity in a single pass with "
   "native parallelisability — critical for large-scale pipelines [10, 11].")

mh("3  State of the Art")
mh("3.1  Post-Quantum Cryptography: CRYSTALS-Kyber / ML-KEM", 2)
mp("FIPS 203 (ML-KEM), finalised August 2024 [12], targets primitives resistant to "
   "Shor's and Grover's algorithms. Kyber is built on Module-LWE over "
   "$R_q = \\mathbb{Z}_q[X]/(X^n + 1)$.")
mh("3.1.1  Mathematical Specification", 3)
mp("**KeyGen:** sample $A \\in R_q^{k \\times k}$; $\\mathbf{s}, \\mathbf{e}$ with small "
   "coefficients; compute $\\mathbf{b} = A\\mathbf{s} + \\mathbf{e}$. "
   "Public key $pk = (A, \\mathbf{b})$; secret key $sk = \\mathbf{s}$.\n\n"
   "**Encaps($pk$):** sample $\\mathbf{r}, \\mathbf{e}_1, e_2$; compute\n"
   "$$\\mathbf{u} = A^\\top \\mathbf{r} + \\mathbf{e}_1, \\quad "
   "v = \\mathbf{b}^\\top \\mathbf{r} + e_2 + \\lfloor q/2 \\rfloor m$$\n"
   "ciphertext $ct = (\\mathbf{u}, v)$; shared secret $K = \\mathrm{KDF}(m)$.\n\n"
   "**Decaps($sk, ct$):** recover "
   "$m' = \\lfloor (2/q)(v - \\mathbf{s}^\\top \\mathbf{u}) \\rceil \\bmod 2$; "
   "derive $K = \\mathrm{KDF}(m')$. "
   "The Fujisaki–Okamoto transform upgrades IND-CPA to IND-CCA2.")

mh("3.2.2  Variational Autoencoders", 3)
mp("A VAE is trained to maximise the Evidence Lower Bound (ELBO) [20]:\n\n"
   "$$\\mathcal{L}(\\theta,\\phi;x) = \\mathbb{E}_{q_\\phi(z|x)}[\\log p_\\theta(x|z)] "
   "- \\mathrm{KL}(q_\\phi(z|x) \\| p(z)) \\tag{3}$$\n\n"
   "The KL term regularises the latent space toward a standard Gaussian, enabling "
   "both compression and generative sampling.")

mh("3.3  Comparative Analysis", 2)
mtbl(
    ["Reference","Year","Components","Strengths","Limitations"],
    [
        ["[22]","2024","FHE + feature extraction","Encrypted computation","No PQ, 120× overhead"],
        ["[23]","2024","VCS + RNS","Realistic leakage model","Depends on RNS quality"],
        ["[24]","2024","AI + intrusion detection","Big data cybersecurity","No PQ primitives"],
        ["[25]","2024","FHE + access control","Massive data protection","No VAE, limited scalability"],
        ["[14]","2024","Blockchain + hybrid PQ sig","Quantum-resistant signature","Healthcare-specific"],
        ["**This work**","2025","VAE + AES-256-GCM + ML-KEM-1024",
         "Joint pipeline; reproducible; PQ + learned compression",
         "CPU-only shim; VAE gains on images only"],
    ]
)

mh("4  Methodology")
mh("4.1  System Pipeline Architecture", 2)
mp("The pipeline applies three sequential transformations:\n\n"
   "> $D \\xrightarrow{\\text{VAE}} z \\xrightarrow{\\text{AES-256-GCM}} c "
   "\\xrightarrow{\\text{ML-KEM}} (c,\\, K_{\\mathrm{enc}})$")
mfig("fig01_pipeline_architecture.png",
     "Fig. 1  Architecture of the proposed hybrid security pipeline.")

mp("**Critical limitation:** The AES-GCM tag guarantees integrity of $z$, not of $D$. "
   "Because the VAE is lossy, $\\hat{D} \\neq D$ in general. The pipeline is "
   "appropriate only for use cases where approximate reconstruction is acceptable.")

mh("4.2  Threat Model", 2)
mtbl(
    ["Target","Vulnerabilities","Attacks","Countermeasures"],
    [
        ["AES-256-GCM","Timing leaks; nonce reuse","DPA, CPA, forgery","Constant-time; random nonces"],
        ["ML-KEM/Kyber","Decapsulation side-channel","FO oracle, SCA","FO transform; masked variant"],
        ["VAE","Model inversion; adversarial","FGSM, PGD, C&W","Latent z encrypted by AES-GCM"],
    ]
)

mh("4.2.2  Algorithm 1: HybridEncrypt", 3)
A("""
```
Algorithm 1. HybridEncrypt(D, pk_Kyber)
Input:  D (data blob),  pk_Kyber (recipient public key)

1.  z        ← VAE_Encode(D)                   # Layer 1: compress
2.  K        ←ᴿ {0,1}^256                      # uniform AES session key
3.  N        ←ᴿ {0,1}^96                       # 96-bit nonce
4.  (ct, τ)  ← AES-GCM_K(z, N)                # Layer 2: encrypt + authenticate
5.  (K_enc, _) ← ML-KEM.Encaps(pk_Kyber)      # Layer 3: encapsulate K
6.  Return (ct, K_enc, N, τ)

HybridDecrypt(ct, K_enc, N, τ, sk_Kyber):
1.  K   ← ML-KEM.Decaps(sk_Kyber, K_enc)
2.  z   ← AES-GCM⁻¹_K(ct, N, τ)              # abort if tag FAIL
3.  D̂  ← VAE_Decode(z)
4.  Return D̂
```
""")

mh("4.3  Experimental Environment", 2)
mtbl(
    ["Element","Details"],
    [
        ["System & OS","Windows 11 Pro, Python 3.13.7, single-node"],
        ["Symmetric crypto","cryptography.hazmat AES-256-GCM (NIST FIPS 197)"],
        ["Post-quantum KEM","kyber-py pure-Python shim; ML-KEM-1024 (FIPS 203)"],
        ["Dataset D1","3.79 GB, 126 files: 75 JPEG (0.49 GB), 50 MP4 (3.30 GB), 1 JSON"],
        ["Measurement","time.perf_counter(); 10 reps/point; median [IQR] reported"],
    ]
)

mh("4.3.2  VAE Architecture", 3)
mtbl(
    ["Component","Architecture","Hyperparameters"],
    [
        ["Encoder","Input → Conv2D(32) → Conv2D(64) → FC(512) → [µ, log σ²]","ReLU; BatchNorm"],
        ["Latent space","z ~ 𝒩(µ, σ²I), dim k = 128","6× CR on 32×32 input"],
        ["Decoder","z → FC(512) → ConvTranspose2D(64) → Output","Sigmoid; MSE + β·KL"],
        ["Training (30 epochs, β=0.01)","Adam lr=5×10⁻⁴; β=0.01; 30 epochs; 80/20 split","Best val. PSNR: 14.85 dB (ep.28); KL: 0.05–0.14; 75 images; collapse resolved"],
    ]
)

mh("4.4  KEM–DEM Security Composition", 2)
mp("Under H1–H3 (MLWE hardness, AES-GCM INT-CTXT, HKDF domain separation), "
   "the KEM-DEM composition inherits IND-CCA2 [27]. "
   "This argument is informal — a formal proof in the random oracle model is "
   "left for future work.")

mh("5  Experimental Results")
mh("5.1  Dataset and Protocol", 2)
mtbl(
    ["File type","Count","Total (GB)","Mean (MB)","Fraction (%)"],
    [
        ["Binary (MP4)","50","3.30","67.3","87.1"],
        ["Images (JPEG)","75","0.49","6.6","12.9"],
        ["Text (JSON)","1","<0.01","0.1","<0.1"],
        ["**Total D1**","**126**","**3.79**","**30.9**","**100.0**"],
    ]
)

mh("5.2  Compression and Quality", 2)
mfig("fig03_compression_ratios.png",
     "Fig. 5  Compression ratios on D1. Classical lossless compressors achieve ≈1× "
     "on JPEG/MP4 (already compressed). VAE achieves 6× on resized 32×32 input.")
mfig("fig_pareto_frontier.png",
     "Fig. 6  Multi-objective trade-off: compression ratio vs. PSNR vs. throughput.")
mfig("fig02_vae_training_history.png",
     "Fig. 7  VAE training history (30 epochs, β=0.01): KL rose from ≈10⁻⁵ to 0.05–0.14; "
     "best val. PSNR = 14.85 dB at epoch 28. Posterior collapse resolved.")
mfig("fig07_vae_reconstruction.png",
     "Fig. 8  VAE reconstructions (30 epochs, β=0.01). Top: six 32×32 test patterns. "
     "Bottom: reconstructions with per-pattern SSIM scores "
     "(gradient=0.464, circle=0.044, noise=0.011, stripes=0.119, checkerboard=0.019, gray=0.945). "
     "All outputs are content-specific — posterior collapse resolved.")

mh("5.3  Cryptographic Costs and Latency", 2)
mfig("fig06_latency_breakdown.png",
     "Fig. 9  Per-layer latency breakdown (Pipeline B). "
     "AES + KEM dominate measured crypto-only time.")

mtbl(
    ["File type","L_AES (ms)","L_KEM (ms)","L_total (ms)","Crypto (% measured)"],
    [
        ["Images (JPEG)","6.601","24.833","31.370 [±3.5]","100.0"],
        ["Binary (MP4)","40.270","29.632","71.645 [±79.6]","97.6"],
        ["**All D1**","**7.126**","**26.496**","**34.235 [±31.0]**","**98.2**"],
    ]
)
A("*Table 6. Median per-layer latency by file type, Pipeline B.*\n")

mfig("fig06_aes_microbench.png",
     "Fig. 10  AES-256-GCM throughput: 109 MB/s (1 KB) → 1,387 MB/s (64 KB).")
mtbl(
    ["Size","Enc. throughput (MB/s)","Dec. throughput (MB/s)"],
    [
        ["1 KB","109","112"],
        ["4 KB","532","537"],
        ["16 KB","997","1,027"],
        ["64 KB","1,387","1,344"],
        ["256 KB","1,511","1,510"],
        ["1 MB","792","912"],
    ]
)
A("*Table 6b. AES-256-GCM throughput by plaintext size.*\n")

mtbl(
    ["Algorithm","KeyGen (µs)","Encaps (µs)","Decaps (µs)"],
    [
        ["Kyber-512","3,827","5,297","7,661"],
        ["Kyber-768","7,965","9,927","13,914"],
        ["**Kyber-1024 ★**","**12,590**","**14,924**","**19,897**"],
        ["ML-KEM-512","4,709","6,581","9,481"],
        ["ML-KEM-768","8,167","10,463","14,504"],
        ["**ML-KEM-1024 ★**","**12,080**","**14,874**","**19,455**"],
    ]
)
A("*Table 7. ML-KEM latency (pure-Python shim). ★ = used in this study.*\n")

mh("5.4  Security Validation", 2)
mfig("fig09_security_validation.png",
     "Fig. 11  NIST CAVP validation: 775/775 test vectors passed (100%).")
mtbl(
    ["Test suite","Total","Pass","Pass rate"],
    [
        ["AES-256-GCM Encrypt KAT","375","375","100.0%"],
        ["AES-256-GCM Decrypt (incl. FAIL)","300","300","100.0%"],
        ["ML-KEM-1024 Keygen/Encaps/Decaps","100","100","100.0%"],
        ["**Total**","**775**","**775**","**100.0%**"],
    ]
)
A("*Table 8. NIST CAVP validation results.*\n")

mfig("fig10a_kem_latency_vs_supercop.png",
     "Fig. 12a  ML-KEM latency: pure-Python shim vs. SUPERCOP bare-metal reference.")
mfig("fig10b_kem_sizes_vs_nist.png",
     "Fig. 12b  ML-KEM parameter sizes vs. NIST KAT reference values (exact match).")

mh("5.5  Ablation Study", 2)
mfig("fig08_ablation.png",
     "Fig. 13  Ablation study: throughput by configuration and file type.")
mtbl(
    ["Config","Img (MB/s)","Bin (MB/s)","Avg (MB/s)","PQ-hardened","Compression"],
    [
        ["A: zstd+AES+KEM","121.7","220.8","114.2","Yes","zstd"],
        ["**B: raw+AES+KEM**","**203.8**","**578.3**","**260.7**","**Yes**","**None**"],
        ["C: raw+AES","993.6","912.8","635.5","No","None"],
        ["D: gzip+AES+KEM","23.9","24.6","16.1","Yes","gzip (slow on entropy)"],
    ]
)
A("*Table 9. Ablation study. B = recommended baseline for PQ security.*\n")

mh("5.6  Scalability", 2)
mfig("fig09_scalability_scatter.png",
     "Fig. 14  Scalability on D1: $\\beta = -0.09$ MB/s/GB, $R^2 < 0.01$ (flat).")
mtbl(
    ["Data (GB)","Throughput (MB/s)","Latency (s)","Files"],
    [
        ["0.60","819.4","0.75","53"],
        ["1.16","1,056.3","0.95","71"],
        ["1.66","1,014.6","1.27","78"],
        ["2.34","949.6","1.82","90"],
        ["2.87","834.9","2.46","108"],
        ["3.30","961.4","2.34","111"],
        ["3.79","954.2","2.58","126"],
    ]
)
A("*Table 10. Scalability results on D1. Mean 941.5 MB/s, essentially flat.*\n")
mfig("fig10_cumulative_latency.png",
     "Fig. 15  Throughput stability vs. cumulative data size on D1.")

mh("6  Discussion")
mp("The ML-KEM/Kyber-1024 overhead of 47.4 ms per key exchange (pure-Python shim) "
   "is the dominant per-file cost for small files, but amortises for large binaries. "
   "On bare-metal native liboqs, this reduces to ~50–100 µs (SUPERCOP).")
mp("The VAE (30 epochs, β=0.01, best PSNR = 14.85 dB) confirms both the compression "
   "potential and the key limitation of the image-trained model: smooth inputs "
   "reconstruct well (gradient SSIM = 0.464, uniform gray SSIM = 0.945), while "
   "high-frequency patterns and binary/tabular data show near-zero compression gain. "
   "An entropy-aware file-type router that bypasses the VAE for high-entropy content "
   "is the correct engineering remedy. Critically, posterior collapse — where the "
   "encoder ignores all inputs and outputs a constant gray — was resolved by "
   "reducing β from 1.0 to 0.01, allowing the reconstruction loss to establish "
   "structure before KL regularisation engages.")
mp("The originality of this work lies not in any individual component, but in "
   "their vertical integration and joint quantitative evaluation on real "
   "heterogeneous data, with explicit cost characterisation of each layer.")

mh("Limitations", 2)
for lim in [
    "(i) CPU-only pure-Python implementation — ML-KEM latency figures are upper "
    "bounds; native liboqs reduces this ~200×.",
    "(ii) Experimental scale (3.79 GB) is below industrial volumes (petabytes); "
    "large-scale (>30 GB) scalability uncharacterised.",
    "(iii) VAE compression specific to images — binary/tabular show near-zero gain.",
    "(iv) KEM-DEM composition argument is informal — formal RO-model proof is "
    "future work.",
]:
    A(f"- {lim}")
A("")

mh("7  Conclusion")
mp("This work presents and evaluates a modular hybrid security pipeline integrating "
   "VAE compression, AES-256-GCM, and ML-KEM/Kyber-1024 on 3.79 GB of real "
   "heterogeneous data (126 files: 75 JPEG, 50 MP4, 1 JSON).")
mp("**Key findings:**")
for kf in [
    "NIST CAVP: 775/775 = 100% pass rate for AES-256-GCM and ML-KEM-1024.",
    "ML-KEM overhead: 47.4 ms/exchange (pure-Python); ~50–100 µs on native bare-metal.",
    "Classical compression fails on high-entropy data: gzip collapses throughput to 24 MB/s.",
    "No throughput collapse: mean 941.5 MB/s, $\\beta = -0.09$ MB/s/GB, $R^2 < 0.01$.",
    "VAE (30 epochs, β=0.01): best PSNR = 14.85 dB; posterior collapse resolved "
    "(KL ≈ 0.05–0.14); smooth inputs SSIM up to 0.945; high-frequency patterns remain "
    "challenging (SSIM 0.011–0.019) due to training set size (68 images).",
]:
    A(f"1. {kf}")
A("")
mp("Future work: (i) bare-metal Linux with native liboqs; (ii) formal KEM-DEM proof; "
   "(iii) entropy-aware file-type router; (iv) PSNR > 25 dB VAE; (v) >30 GB scalability; "
   "(vi) extension to IoT/streaming data.")

mh("Declarations")
for label, text in [
    ("**Funding:**", "No specific grant from any funding agency."),
    ("**Competing interests:**", "The authors declare no competing interests."),
    ("**AI use disclosure:**", "Claude (Anthropic) used for English editing and code support. "
     "All scientific content is the sole responsibility of the human authors."),
    ("**Ethics:**", "Not applicable — no human subjects or personal data."),
]:
    A(f"\n{label} {text}")
A("")

mh("References")
for ref in refs:
    A(f"\n{ref}")
A("")

# Write markdown
OUT_M.write_text("\n".join(md), encoding="utf-8")
print(f"[MD]   Saved: {OUT_M}")
print("Done.")
